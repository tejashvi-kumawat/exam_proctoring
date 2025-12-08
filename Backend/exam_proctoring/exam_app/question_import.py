# backend/exam_proctoring/exam_app/question_import.py
"""
Bulk question import from JSON and CSV files
Supports multiple question formats and validation
Includes image support via base64 encoding or URLs
"""
import json
import csv
import io
import base64
import os
from urllib.parse import urlparse
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
from django.db import transaction
from django.core.exceptions import ValidationError
from .models import Exam, Question, Option


class QuestionImportError(Exception):
    """Custom exception for question import errors"""
    pass


def decode_base64_image(base64_string, filename='image.png'):
    """
    Decode base64 string to Django ContentFile
    
    Supports formats:
    - "data:image/png;base64,iVBORw0KGgo..." (with data URI)
    - "iVBORw0KGgo..." (raw base64)
    """
    try:
        # Remove data URI prefix if present
        if ',' in base64_string:
            header, data = base64_string.split(',', 1)
            # Extract format from header (e.g., "data:image/png;base64")
            if 'image/' in header:
                ext = header.split('image/')[1].split(';')[0]
                filename = f'image.{ext}'
        else:
            data = base64_string
        
        # Decode base64
        image_data = base64.b64decode(data)
        return ContentFile(image_data, name=filename)
    except Exception as e:
        raise QuestionImportError(f"Invalid base64 image: {str(e)}")


def get_image_from_url_or_path(image_path, exam_id=None):
    """
    Get image from URL or local path
    
    Returns ContentFile or None
    """
    if not image_path or not image_path.strip():
        return None
    
    image_path = image_path.strip()
    
    # Check if it's a URL
    parsed = urlparse(image_path)
    if parsed.scheme in ['http', 'https']:
        # It's a URL - download it
        try:
            import requests
            response = requests.get(image_path, timeout=10)
            response.raise_for_status()
            
            # Get filename from URL or use default
            filename = os.path.basename(parsed.path) or 'image.png'
            if not filename.endswith(('.png', '.jpg', '.jpeg', '.gif', '.webp')):
                filename = 'image.png'
            
            return ContentFile(response.content, name=filename)
        except Exception as e:
            raise QuestionImportError(f"Failed to download image from URL {image_path}: {str(e)}")
    else:
        # It's a local path - read file
        try:
            if os.path.exists(image_path):
                with open(image_path, 'rb') as f:
                    filename = os.path.basename(image_path)
                    return ContentFile(f.read(), name=filename)
            else:
                raise QuestionImportError(f"Image file not found: {image_path}")
        except Exception as e:
            raise QuestionImportError(f"Failed to read image file {image_path}: {str(e)}")


def validate_question_data(question_data, question_index):
    """Validate a single question data structure"""
    errors = []
    
    if 'question_text' not in question_data or not question_data['question_text']:
        errors.append(f"Question {question_index + 1}: Missing or empty 'question_text'")
    
    question_type = question_data.get('question_type', 'MCQ')
    if question_type not in ['MCQ', 'TF', 'SA', 'TEXT', 'IMAGE_UPLOAD']:
        errors.append(f"Question {question_index + 1}: Invalid question_type. Must be MCQ, TF, SA, TEXT, or IMAGE_UPLOAD")
    
    if question_type in ['MCQ', 'TF']:
        if 'options' not in question_data or not isinstance(question_data['options'], list):
            errors.append(f"Question {question_index + 1}: Missing or invalid 'options' array")
        elif len(question_data['options']) < 2:
            errors.append(f"Question {question_index + 1}: Must have at least 2 options")
        else:
            # Check if at least one option is marked as correct
            correct_options = [opt for opt in question_data['options'] if opt.get('is_correct', False)]
            if len(correct_options) == 0:
                errors.append(f"Question {question_index + 1}: At least one option must be marked as correct")
            elif len(correct_options) > 1 and question_type == 'TF':
                errors.append(f"Question {question_index + 1}: True/False questions must have exactly one correct option")
    
    # For SA, TEXT, IMAGE_UPLOAD - options are optional
    if question_type in ['SA', 'TEXT', 'IMAGE_UPLOAD']:
        # Options can be empty or not provided
        pass
    
    marks = question_data.get('marks', 1)
    if not isinstance(marks, (int, float)) or marks <= 0:
        errors.append(f"Question {question_index + 1}: 'marks' must be a positive number")
    
    return errors


def import_questions_from_json(exam_id, json_data, overwrite=False):
    """
    Import questions from JSON data
    
    Expected JSON format:
    {
        "questions": [
            {
                "question_text": "What is 2+2?",
                "question_type": "MCQ",
                "marks": 1,
                "order": 1,
                "options": [
                    {"option_text": "3", "is_correct": false, "order": 0},
                    {"option_text": "4", "is_correct": true, "order": 1},
                    {"option_text": "5", "is_correct": false, "order": 2}
                ]
            },
            {
                "question_text": "Python is a programming language",
                "question_type": "TF",
                "marks": 1,
                "order": 2,
                "options": [
                    {"option_text": "True", "is_correct": true, "order": 0},
                    {"option_text": "False", "is_correct": false, "order": 1}
                ]
            }
        ]
    }
    
    Alternative format (simpler):
    [
        {
            "question_text": "...",
            "options": ["Option 1", "Option 2", "Option 3"],
            "correct_option": 1  # index of correct option
        }
    ]
    """
    try:
        exam = Exam.objects.get(id=exam_id)
    except Exam.DoesNotExist:
        raise QuestionImportError(f"Exam with id {exam_id} does not exist")
    
    # Parse JSON if string
    if isinstance(json_data, str):
        try:
            data = json.loads(json_data)
        except json.JSONDecodeError as e:
            raise QuestionImportError(f"Invalid JSON format: {str(e)}")
    else:
        data = json_data
    
    # Handle different JSON formats
    if isinstance(data, list):
        # Simple format: list of questions
        questions_data = data
    elif isinstance(data, dict) and 'questions' in data:
        # Standard format: {"questions": [...]}
        questions_data = data['questions']
    else:
        raise QuestionImportError("Invalid JSON structure. Expected array of questions or object with 'questions' key")
    
    if not isinstance(questions_data, list):
        raise QuestionImportError("Questions must be an array")
    
    if len(questions_data) == 0:
        raise QuestionImportError("No questions found in JSON data")
    
    # Validate all questions first
    all_errors = []
    for i, question_data in enumerate(questions_data):
        errors = validate_question_data(question_data, i)
        all_errors.extend(errors)
    
    if all_errors:
        raise QuestionImportError("Validation errors:\n" + "\n".join(all_errors))
    
    # Import questions
    imported_count = 0
    errors = []
    
    with transaction.atomic():
        if overwrite:
            # Delete existing questions
            Question.objects.filter(exam=exam).delete()
        
        for i, question_data in enumerate(questions_data):
            try:
                # Handle simple format
                if 'options' in question_data and isinstance(question_data['options'], list):
                    if len(question_data['options']) > 0 and isinstance(question_data['options'][0], str):
                        # Simple format: ["Option 1", "Option 2"]
                        options_list = question_data['options']
                        correct_index = question_data.get('correct_option', 0)
                        
                        # Convert to standard format
                        question_data['options'] = [
                            {
                                'option_text': opt,
                                'is_correct': idx == correct_index,
                                'order': idx
                            }
                            for idx, opt in enumerate(options_list)
                        ]
                
                # Create question
                question_type = question_data.get('question_type', 'MCQ')
                
                # Handle question image
                question_image = None
                if 'question_image' in question_data and question_data['question_image']:
                    image_data = question_data['question_image']
                    if isinstance(image_data, str):
                        if image_data.startswith('data:image') or len(image_data) > 100:
                            # Base64 encoded image
                            question_image = decode_base64_image(image_data, f'question_{i+1}.png')
                        else:
                            # URL or path
                            question_image = get_image_from_url_or_path(image_data, exam_id)
                
                question = Question.objects.create(
                    exam=exam,
                    question_text=question_data['question_text'],
                    question_type=question_type,
                    marks=question_data.get('marks', 1),
                    order=question_data.get('order', i + 1),
                    question_image=question_image
                )
                
                # Create options (only for MCQ/TF)
                if question_type in ['MCQ', 'TF']:
                    options_data = question_data.get('options', [])
                    for opt_idx, opt_data in enumerate(options_data):
                        option_text = ''
                        option_image = None
                        
                        if isinstance(opt_data, dict):
                            option_text = opt_data.get('option_text', '')
                            
                            # Handle option image
                            if 'option_image' in opt_data and opt_data['option_image']:
                                image_data = opt_data['option_image']
                                if isinstance(image_data, str):
                                    if image_data.startswith('data:image') or len(image_data) > 100:
                                        # Base64 encoded image
                                        option_image = decode_base64_image(image_data, f'option_{i+1}_{opt_idx}.png')
                                    else:
                                        # URL or path
                                        option_image = get_image_from_url_or_path(image_data, exam_id)
                            
                            Option.objects.create(
                                question=question,
                                option_text=option_text,
                                option_image=option_image,
                                is_correct=opt_data.get('is_correct', False),
                                order=opt_data.get('order', opt_idx)
                            )
                        else:
                            # Handle string options
                            Option.objects.create(
                                question=question,
                                option_text=str(opt_data),
                                is_correct=False,
                                order=opt_idx
                            )
                
                imported_count += 1
                
            except Exception as e:
                errors.append(f"Question {i + 1}: {str(e)}")
        
        if errors and imported_count == 0:
            raise QuestionImportError("Failed to import any questions:\n" + "\n".join(errors))
    
    # Recalculate exam total marks
    exam.save()  # This will trigger auto_calculate_total if enabled
    
    return {
        'imported': imported_count,
        'total': len(questions_data),
        'errors': errors
    }


def import_questions_from_csv(exam_id, csv_content, overwrite=False):
    """
    Import questions from CSV data
    
    CSV Format:
    - Column 1: Question Text
    - Column 2: Question Type (MCQ, TF, SA, TEXT, IMAGE_UPLOAD)
    - Column 3: Marks
    - Column 4+: Options (for MCQ/TF) - last column indicates correct option number (1-based)
    
    Example CSV:
    Question,Type,Marks,Option1,Option2,Option3,Option4,Correct
    "What is 2+2?",MCQ,1,3,4,5,6,2
    "Python is a language",TF,1,True,False,,
    "Explain recursion",SA,5,,,,
    """
    try:
        exam = Exam.objects.get(id=exam_id)
    except Exam.DoesNotExist:
        raise QuestionImportError(f"Exam with id {exam_id} does not exist")
    
    # Parse CSV
    try:
        if isinstance(csv_content, bytes):
            csv_content = csv_content.decode('utf-8')
        
        csv_reader = csv.reader(io.StringIO(csv_content))
        rows = list(csv_reader)
        
        if len(rows) < 2:  # Header + at least one question
            raise QuestionImportError("CSV file must have at least one question row")
        
        # Skip header row
        header = rows[0] if len(rows) > 0 else []
        data_rows = rows[1:]
        
    except Exception as e:
        raise QuestionImportError(f"Error parsing CSV: {str(e)}")
    
    # Convert CSV rows to question data format
    questions_data = []
    for row_idx, row in enumerate(data_rows):
        if not row or len(row) < 3:  # Need at least Question, Type, Marks
            continue
        
        question_text = row[0].strip() if len(row) > 0 else ''
        question_type = row[1].strip().upper() if len(row) > 1 else 'MCQ'
        marks = row[2].strip() if len(row) > 2 else '1'
        
        if not question_text:
            continue
        
        # Parse marks
        try:
            marks = float(marks)
        except ValueError:
            marks = 1.0
        
        question_data = {
            'question_text': question_text,
            'question_type': question_type,
            'marks': marks,
            'order': row_idx + 1
        }
        
        # Handle question image (check if there's a QuestionImage column or in a specific position)
        # Look for columns that might contain image data
        # Format: Question,Type,Marks,QuestionImage,Option1,Option2,...
        question_image_col = None
        if len(header) > 3:
            # Check if there's a "QuestionImage" column
            for col_idx, col_name in enumerate(header):
                if col_name.lower() in ['questionimage', 'question_image', 'qimage', 'q_image']:
                    if len(row) > col_idx:
                        question_image_col = row[col_idx].strip() if row[col_idx] else None
                    break
        
        if question_image_col:
            question_data['question_image'] = question_image_col
        
        # Handle options for MCQ/TF
        if question_type in ['MCQ', 'TF']:
            options = []
            # Get options from columns 3 onwards (skip Question, Type, Marks, and QuestionImage if present)
            start_col = 4 if question_image_col else 3
            option_cols = row[start_col:] if len(row) > start_col else []
            
            # Find correct option index (last column or explicit "Correct" column)
            correct_index = None
            if len(option_cols) > 0:
                # Check if last column is a number (correct option index)
                try:
                    correct_index = int(option_cols[-1]) - 1  # Convert to 0-based
                    option_cols = option_cols[:-1]  # Remove correct index column
                except ValueError:
                    # Last column is not a number, might be an option itself
                    pass
            
            # Filter out empty options and separate text from images
            # Format: Option1,Option1Image,Option2,Option2Image,...
            option_texts = []
            option_images = []
            
            # Simple approach: treat every other column as option text/image pairs
            # Or just treat all as text options
            for opt in option_cols:
                opt = opt.strip() if opt else ''
                if opt:
                    # Check if it looks like base64 or URL
                    if opt.startswith('data:image') or opt.startswith('http') or len(opt) > 100:
                        # It's an image - attach to last option or create image-only option
                        if option_texts:
                            option_images.append(opt)
                        else:
                            # Image-only option
                            option_texts.append('')
                            option_images.append(opt)
                    else:
                        # It's text
                        option_texts.append(opt)
                        option_images.append(None)
            
            # Ensure we have at least 2 options
            if len(option_texts) < 2:
                if len(option_cols) >= 2:
                    # Re-process: treat all as text options
                    option_texts = [opt.strip() for opt in option_cols if opt.strip()]
                    option_images = [None] * len(option_texts)
                else:
                    raise QuestionImportError(f"Row {row_idx + 2}: MCQ/TF questions need at least 2 options")
            
            # If correct_index not found, default to first option
            if correct_index is None:
                correct_index = 0
            
            # Create options array
            for idx, opt_text in enumerate(option_texts):
                opt_data = {
                    'option_text': opt_text,
                    'is_correct': idx == correct_index,
                    'order': idx
                }
                # Add image if available
                if idx < len(option_images) and option_images[idx]:
                    opt_data['option_image'] = option_images[idx]
                question_data.setdefault('options', []).append(opt_data)
        
        questions_data.append(question_data)
    
    if len(questions_data) == 0:
        raise QuestionImportError("No valid questions found in CSV file")
    
    # Use existing JSON import function
    return import_questions_from_json(exam_id, questions_data, overwrite)


def extract_image_from_docx_paragraph(para):
    """Extract image from a DOCX paragraph"""
    try:
        from docx.document import Document as _Document
        from docx.oxml.text.paragraph import CT_P
        from docx.oxml.ns import qn
        
        # Check for inline shapes (images)
        for run in para.runs:
            for drawing in run.element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
                # Get image relationship ID
                embed_id = drawing.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if embed_id:
                    return embed_id
        return None
    except:
        return None


def get_image_from_docx_relationship(doc, embed_id):
    """Get image content from DOCX relationship"""
    try:
        # Access the document's part relationships
        part = doc.part
        if hasattr(part, 'related_parts'):
            for rel_id, rel_part in part.related_parts.items():
                if rel_id == embed_id:
                    return rel_part.blob
        return None
    except:
        return None


def import_questions_from_docx(exam_id, docx_content, overwrite=False):
    """
    Import questions from DOCX (Word) file with image support
    
    Expected DOCX format:
    - Questions should be numbered (1., 2., etc.) or marked with Q:
    - Options should be marked with a), b), c), d) or A), B), C), D)
    - Correct answer can be marked with [Correct] or * at the end
    - Marks can be specified in brackets like (2 marks) or [2 marks]
    - Images embedded in document will be automatically extracted
    - Question type can be inferred or specified
    
    Example:
    Q1. What is 2+2? (1 mark)
    [Image embedded here]
    a) 3
    b) 4 [Correct]
    c) 5
    d) 6
    """
    try:
        exam = Exam.objects.get(id=exam_id)
    except Exam.DoesNotExist:
        raise QuestionImportError(f"Exam with id {exam_id} does not exist")
    
    try:
        # Try to import docx - check if it's available
        try:
            from docx import Document
            from docx.oxml.ns import qn
        except ImportError as import_err:
            import sys
            python_path = sys.executable
            error_msg = f"python-docx library is not installed in the current Python environment.\n"
            error_msg += f"Current Python: {python_path}\n"
            error_msg += f"Import error: {str(import_err)}\n\n"
            error_msg += "To fix this, run:\n"
            error_msg += f"  {python_path} -m pip install python-docx==1.1.2\n"
            error_msg += "Or activate your virtual environment and run:\n"
            error_msg += "  pip install python-docx==1.1.2"
            raise QuestionImportError(error_msg)
        
        import io
        import re
        
        # Parse DOCX
        if isinstance(docx_content, bytes):
            doc = Document(io.BytesIO(docx_content))
            docx_bytes = docx_content
        else:
            docx_bytes = docx_content.read()
            doc = Document(io.BytesIO(docx_bytes))
        
        # Extract all images from document
        images_dict = {}
        try:
            # Access document relationships to get images
            part = doc.part
            if hasattr(part, 'related_parts'):
                for rel_id, rel_part in part.related_parts.items():
                    if hasattr(rel_part, 'blob'):
                        # Determine image format from content type
                        content_type = getattr(rel_part, 'content_type', '')
                        ext = 'png'
                        if 'jpeg' in content_type or 'jpg' in content_type:
                            ext = 'jpg'
                        elif 'gif' in content_type:
                            ext = 'gif'
                        elif 'webp' in content_type:
                            ext = 'webp'
                        
                        images_dict[rel_id] = {
                            'content': rel_part.blob,
                            'ext': ext
                        }
        except Exception as e:
            print(f"Warning: Could not extract all images: {e}")
        
        questions_data = []
        current_question = None
        current_options = []
        question_number = 0
        current_question_image = None
        current_option_image = None
        
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            
            # Check for images in this paragraph
            para_image = None
            try:
                # Check for inline shapes (images) in paragraph
                for run in para.runs:
                    if run._element.xml:
                        # Look for image relationships
                        import xml.etree.ElementTree as ET
                        root = ET.fromstring(run._element.xml)
                        for blip in root.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
                            embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            if embed_id and embed_id in images_dict:
                                img_data = images_dict[embed_id]
                                # Convert to base64
                                import base64
                                base64_img = base64.b64encode(img_data['content']).decode('utf-8')
                                para_image = f"data:image/{img_data['ext']};base64,{base64_img}"
                                break
            except Exception as e:
                print(f"Warning: Could not extract image from paragraph {para_idx}: {e}")
            
            if not text and not para_image:
                continue
            
            # Check if it's a question (starts with Q, number, or question mark)
            if text and ((text[0].isdigit() and ('.' in text[:5] or ')' in text[:5])) or \
               text.lower().startswith('q') or \
               '?' in text[:100]):
                
                # Save previous question if exists
                if current_question:
                    question_type = current_question.get('question_type', 'MCQ')
                    
                    # If MCQ/TF but no options found, change to SA type
                    if question_type in ['MCQ', 'TF']:
                        if len(current_options) >= 2:
                            current_question['options'] = current_options
                        else:
                            # No options found - change to SA (Short Answer) type
                            current_question['question_type'] = 'SA'
                    
                    questions_data.append(current_question)
                    current_options = []
                    current_option_image = None
                
                # Extract question text and marks
                question_text = text
                marks = 1
                question_type = 'MCQ'
                
                # Extract marks from text (e.g., "(2 marks)", "[2 marks]", "(2)")
                marks_match = re.search(r'\((\d+)\s*marks?\)|\[(\d+)\s*marks?\]|\((\d+)\)', text, re.IGNORECASE)
                if marks_match:
                    marks = int(marks_match.group(1) or marks_match.group(2) or marks_match.group(3))
                    question_text = re.sub(r'\((\d+)\s*marks?\)|\[(\d+)\s*marks?\]|\((\d+)\)', '', question_text, flags=re.IGNORECASE).strip()
                
                # Detect question type
                if 'true' in text.lower() and 'false' in text.lower():
                    question_type = 'TF'
                elif 'explain' in text.lower() or 'describe' in text.lower() or 'write' in text.lower():
                    question_type = 'SA'
                elif 'essay' in text.lower() or 'detailed' in text.lower():
                    question_type = 'TEXT'
                elif 'upload' in text.lower() or 'image' in text.lower():
                    question_type = 'IMAGE_UPLOAD'
                
                # Remove question number prefix
                question_text = re.sub(r'^Q\d+[.:]\s*|^\d+[.:)]\s*', '', question_text, flags=re.IGNORECASE).strip()
                
                question_number += 1
                current_question = {
                    'question_text': question_text,
                    'question_type': question_type,
                    'marks': marks,
                    'order': question_number
                }
                
                # Set question image if found
                if para_image:
                    current_question_image = para_image
                    current_question['question_image'] = para_image
                else:
                    current_question_image = None
            
            # Check if it's an option (starts with a), b), c), d) or A), B), C), D)
            # Also check for numbered options like 1), 2), 3), 4)
            elif current_question and text and (
                re.match(r'^[a-dA-D][).]\s+', text) or 
                re.match(r'^[ivx]+[).]\s+', text, re.IGNORECASE) or
                re.match(r'^[1-4][).]\s+', text)
            ):
                # Remove option prefix (a), b), 1), etc.)
                option_text = re.sub(r'^[a-dA-D][).]\s+|^[ivx]+[).]\s+|^[1-4][).]\s+', '', text, flags=re.IGNORECASE).strip()
                
                # Check if it's marked as correct
                is_correct = False
                if '[correct]' in option_text.lower() or '*correct*' in option_text.lower() or option_text.endswith('*'):
                    is_correct = True
                    option_text = re.sub(r'\[correct\]|\*correct\*|\*$', '', option_text, flags=re.IGNORECASE).strip()
                
                # Remove asterisk or other markers
                option_text = option_text.rstrip('*').strip()
                
                if option_text or para_image:
                    opt_data = {
                        'option_text': option_text,
                        'is_correct': is_correct,
                        'order': len(current_options)
                    }
                    
                    # Add image if found
                    if para_image:
                        opt_data['option_image'] = para_image
                        current_option_image = para_image
                    elif current_option_image:
                        # Use previous image if current option doesn't have one
                        opt_data['option_image'] = current_option_image
                    
                    current_options.append(opt_data)
                    current_option_image = None  # Reset after use
            
            # Check if paragraph has only image (no text) - attach to current question/option
            elif para_image:
                if current_question:
                    if current_question.get('question_type') in ['MCQ', 'TF']:
                        # If we're in options mode, attach to last option
                        if current_options:
                            if 'option_image' not in current_options[-1]:
                                current_options[-1]['option_image'] = para_image
                        else:
                            # No options yet, attach to question
                            if 'question_image' not in current_question:
                                current_question['question_image'] = para_image
                    else:
                        # For SA/TEXT/IMAGE_UPLOAD, attach to question
                        if 'question_image' not in current_question:
                            current_question['question_image'] = para_image
            
            # If current question exists and no option pattern, might be continuation of question text
            elif current_question and text and current_question.get('question_type') not in ['MCQ', 'TF']:
                # Append to question text for SA, TEXT, IMAGE_UPLOAD
                current_question['question_text'] += ' ' + text
        
        # Save last question
        if current_question:
            question_type = current_question.get('question_type', 'MCQ')
            
            # If MCQ/TF but no options found, change to SA type
            if question_type in ['MCQ', 'TF']:
                if len(current_options) >= 2:
                    current_question['options'] = current_options
                else:
                    # No options found - change to SA (Short Answer) type
                    current_question['question_type'] = 'SA'
                    # Don't add empty options array - SA questions don't need options
            
            questions_data.append(current_question)
        
        if len(questions_data) == 0:
            raise QuestionImportError("No questions found in DOCX file. Please ensure questions are properly formatted.")
        
        # Use existing JSON import function
        return import_questions_from_json(exam_id, questions_data, overwrite)
        
    except ImportError as e:
        import sys
        python_path = sys.executable
        error_msg = f"python-docx library is not installed in the current Python environment.\n"
        error_msg += f"Current Python: {python_path}\n"
        error_msg += f"Import error: {str(e)}\n\n"
        error_msg += "To fix this, run:\n"
        error_msg += f"  {python_path} -m pip install python-docx==1.1.2\n"
        error_msg += "Or activate your virtual environment and run:\n"
        error_msg += "  pip install python-docx==1.1.2"
        raise QuestionImportError(error_msg)
    except Exception as e:
        raise QuestionImportError(f"Error parsing DOCX file: {str(e)}")


def import_questions_from_file(exam_id, file_path, overwrite=False):
    """Import questions from a JSON file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            json_data = json.load(f)
        return import_questions_from_json(exam_id, json_data, overwrite)
    except FileNotFoundError:
        raise QuestionImportError(f"File not found: {file_path}")
    except json.JSONDecodeError as e:
        raise QuestionImportError(f"Invalid JSON in file: {str(e)}")
    except Exception as e:
        raise QuestionImportError(f"Error reading file: {str(e)}")

